#!/usr/bin/env python3
"""
extract_text.py — Extract raw text from a book file (PDF/EPUB/DOCX/TXT/MD) and
split it into heuristically-detected sections/chapters.

This is a starting point for the orchestrating agent, not ground truth: the
agent is expected to look at manifest.json and re-chunk if the detected
boundaries look wrong (too many/few sections, tiny fragments, etc.).

Usage:
    python3 extract_text.py <input_file> --outdir <output_dir>

Output:
    <outdir>/sections/001_<slug>.txt, 002_<slug>.txt, ...
    <outdir>/manifest.json
        [{"index": 1, "title": "...", "path": "...", "word_count": 1234,
          "detection_method": "pdf_bookmark" | "heading_style" |
                              "regex_chapter" | "markdown_header" |
                              "epub_spine" | "single_section"}]

Dependencies (install first):
    pip install --break-system-packages pypdf python-docx ebooklib beautifulsoup4 lxml
"""

import argparse
import json
import re
import sys
from pathlib import Path

CHAPTER_RE = re.compile(
    r"^\s*(chapter|capitolo|cap\.?|part|parte|section|sezione)\s+([0-9ivxlcdm]+)\b",
    re.IGNORECASE,
)
MD_HEADER_RE = re.compile(r"^#{1,3}\s+(.*)")


def slugify(title, max_len=50):
    s = re.sub(r"[^a-zA-Z0-9]+", "-", title.strip().lower()).strip("-")
    return (s or "section")[:max_len]


def word_count(text):
    return len(text.split())


def write_sections(sections, outdir):
    """sections: list of (title, text, detection_method)"""
    sec_dir = outdir / "sections"
    sec_dir.mkdir(parents=True, exist_ok=True)
    manifest = []
    for i, (title, text, method) in enumerate(sections, start=1):
        fname = f"{i:03d}_{slugify(title)}.txt"
        path = sec_dir / fname
        path.write_text(text, encoding="utf-8")
        manifest.append(
            {
                "index": i,
                "title": title,
                "path": str(path),
                "word_count": word_count(text),
                "detection_method": method,
            }
        )
    (outdir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return manifest


def split_by_regex_or_whole(full_text, method_if_split="regex_chapter"):
    lines = full_text.splitlines()
    boundaries = [i for i, line in enumerate(lines) if CHAPTER_RE.match(line)]
    if not boundaries:
        return [("Full text (no chapters detected)", full_text, "single_section")]
    sections = []
    for idx, start in enumerate(boundaries):
        end = boundaries[idx + 1] if idx + 1 < len(boundaries) else len(lines)
        title = lines[start].strip()
        text = "\n".join(lines[start:end]).strip()
        if text:
            sections.append((title, text, method_if_split))
    return sections


def extract_pdf(path):
    import fitz
    doc = fitz.open(str(path))
    outline = doc.get_toc()
    
    # Check if there is any embedded text
    total_text_len = sum(len(page.get_text() or '') for page in doc)
    
    if total_text_len > 100:
        page_texts = [page.get_text() or '' for page in doc]
        if outline:
            sections = []
            for idx, item in enumerate(outline):
                level, title, page_start = item
                page_start -= 1  # Convert to 0-based index
                page_end = outline[idx + 1][2] - 1 if idx + 1 < len(outline) else len(page_texts)
                text = "\n".join(page_texts[page_start:page_end]).strip()
                if text:
                    sections.append((title, text, "pdf_toc"))
            if sections:
                return sections
        full_text = "\n".join(page_texts)
        return split_by_regex_or_whole(full_text)
    
    raise RuntimeError(
        "Image-only/Scanned PDF detected. This format requires OCR. "
        "As per the G10DC skill specification, alembic does not implement raw OCR. "
        "Please delegate the document extraction to the 'scribe' skill for preprocessing, "
        "verification, and repair, and then feed the resulting digital text directly into alembic."
    )


def extract_docx(path):
    import docx

    d = docx.Document(str(path))
    sections = []
    current_title = "Introduzione"
    current_lines = []
    method = "single_section"

    def flush():
        text = "\n".join(current_lines).strip()
        if text:
            sections.append((current_title, text, method))

    for para in d.paragraphs:
        style = (para.style.name if para.style else "") or ""
        if style.lower().startswith("heading"):
            flush()
            current_title = para.text.strip() or current_title
            current_lines = []
            method = "heading_style"
        else:
            current_lines.append(para.text)
    flush()

    if len(sections) <= 1:
        full_text = "\n".join(p.text for p in d.paragraphs)
        return split_by_regex_or_whole(full_text)
    return sections


def extract_epub(path):
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(str(path))
    sections = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "lxml")
        text = soup.get_text("\n").strip()
        if not text:
            continue
        heading = soup.find(["h1", "h2", "h3"])
        title = heading.get_text().strip() if heading else (item.get_name() or "Sezione")
        sections.append((title, text, "epub_spine"))
    if not sections:
        raise ValueError("No readable content found in EPUB.")
    return sections


def extract_txt_or_md(path):
    full_text = path.read_text(encoding="utf-8", errors="ignore")
    lines = full_text.splitlines()
    md_boundaries = [i for i, line in enumerate(lines) if MD_HEADER_RE.match(line)]
    if md_boundaries:
        sections = []
        for idx, start in enumerate(md_boundaries):
            end = md_boundaries[idx + 1] if idx + 1 < len(md_boundaries) else len(lines)
            title = MD_HEADER_RE.match(lines[start]).group(1).strip()
            text = "\n".join(lines[start:end]).strip()
            if text:
                sections.append((title, text, "markdown_header"))
        if sections:
            return sections
    return split_by_regex_or_whole(full_text)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_file")
    parser.add_argument("--outdir", required=True)
    args = parser.parse_args()

    in_path = Path(args.input_file)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    ext = in_path.suffix.lower()
    try:
        if ext == ".pdf":
            sections = extract_pdf(in_path)
        elif ext == ".docx":
            sections = extract_docx(in_path)
        elif ext == ".epub":
            sections = extract_epub(in_path)
        elif ext in (".txt", ".md", ".markdown"):
            sections = extract_txt_or_md(in_path)
        else:
            print(f"❌ Unsupported extension: {ext}", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"❌ Extraction failed: {e}", file=sys.stderr)
        sys.exit(1)

    manifest = write_sections(sections, outdir)
    total_words = sum(m["word_count"] for m in manifest)
    print(f"✅ Extracted {len(manifest)} section(s), {total_words} words total.")
    print(f"   Manifest: {outdir / 'manifest.json'}")
    for m in manifest[:10]:
        print(f"   [{m['index']:03d}] {m['title'][:60]!r} — {m['word_count']} words ({m['detection_method']})")
    if len(manifest) > 10:
        print(f"   ... and {len(manifest) - 10} more (see manifest.json)")


if __name__ == "__main__":
    main()
